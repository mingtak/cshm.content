# -*- coding: utf-8 -*-
from cshm.content import _
from Products.Five.browser import BrowserView
from Products.Five.browser.pagetemplatefile import ViewPageTemplateFile
from plone import api
from mingtak.ECBase.browser.views import SqlObj
import json
import datetime
import math
from docxtpl import DocxTemplate, Listing
from Products.CMFPlone.utils import safe_unicode
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Basic(BrowserView):

    def getAllStudent(self):
        uid = self.context.UID()
        sqlInstance = SqlObj()
        sqlStr = """SELECT *, training_status_code.status FROM `reg_course`, `training_status_code`, education_code WHERE uid = '{}' 
                 and isAlt = 0 and seatNo IS NOT null and reg_course.training_status = training_status_code.id AND
                 reg_course.education_id = education_code.id ORDER BY reg_course.seatNo""".format(uid)
        return sqlInstance.execSql(sqlStr)

    def getChineseBirthday(self, birthday):
        if birthday:
            year = birthday.year - 1911
            month = birthday.month
            day = birthday.day
            return '%s/%s/%s' %(year, month, day)
        else:
            return ''

    def downloadFile(self, title):
        self.request.response.setHeader('Content-Type', 'application/docx')
        self.request.response.setHeader(
            'Content-Disposition',
            'attachment; filename="%s.docx"' %title
        )
        with open('/tmp/temp.docx', 'rb') as file:
            docs = file.read()
            return docs


class DownloadSignatureDoc(Basic):
    def __call__(self):
        request = self.request
        id = request.get('id')
        type = request.get('type')
        sqlInstance = SqlObj()
        sqlStr = """SELECT * FROM signature WHERE id = {}""".format(id)
        result = sqlInstance.execSql(sqlStr)[0]
        obj = dict(result)
        parameter = {}
        sn = obj['year']
        for i in range(4 - len(str(obj['code']))):
            sn = '0' + str(sn)
        parameter['sn'] = safe_unicode(sn)
        parameter['year'] = safe_unicode(obj['date'].year)
        parameter['month'] = safe_unicode(obj['date'].month)
        parameter['day'] = safe_unicode(obj['date'].day)
        parameter['title'] = safe_unicode(obj['title'])
        parameter['detail_1'] = safe_unicode(obj['detail_1'])
        parameter['detail_2'] = safe_unicode(obj['detail_2'])
        parameter['detail_3'] = safe_unicode(obj['detail_3'])
        parameter['detail_4'] = safe_unicode(obj['detail_4'])
        parameter['department'] = safe_unicode(obj['department'])
        parameter['type'] = safe_unicode(obj['type'])
        parameter['category'] = safe_unicode(obj['category'])
        if type == '1':
            parameter['condition'] = safe_unicode('會辦部門行政服務處')

        filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/signature.docx'

        doc = DocxTemplate(filePath)
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile('%s_簽呈' %sn)


class DownloadOfficialDoc(Basic):
    def __call__(self):
        request = self.request
        uid = request.get('uid')
        parameter = {}
        obj = api.content.get(UID=uid)

        title = obj.title
        recipient = obj.recipient
        parameter['title'] = title
        parameter['docDate'] = safe_unicode(self.getChineseBirthday(obj.docDate))
        parameter['recipient'] = safe_unicode(obj.recipient)
        parameter['docSN'] = safe_unicode(obj.docHeader + obj.docSN)
        parameter['detail_1'] = '1. %s' % safe_unicode(obj.detail_1) if obj.detail_1 else ''
        parameter['detail_2'] = '2. %s' % safe_unicode(obj.detail_2) if obj.detail_2 else ''
        parameter['detail_3'] = '3. %s' % safe_unicode(obj.detail_3) if obj.detail_3 else ''
        parameter['detail_4'] = '4. %s' % safe_unicode(obj.detail_4) if obj.detail_4 else ''
        parameter['detail_5'] = '5. %s' % safe_unicode(obj.detail_5) if obj.detail_5 else ''
        parameter['detail_6'] = '6. %s' % safe_unicode(obj.detail_6) if obj.detail_6 else ''
        parameter['detail_7'] = '7. %s' % safe_unicode(obj.detail_7) if obj.detail_7 else ''
        parameter['detail_8'] = '8. %s' % safe_unicode(obj.detail_8) if obj.detail_8 else ''
        parameter['detail_9'] = '9. %s' % safe_unicode(obj.detail_9) if obj.detail_9 else ''
        parameter['detail_10'] = '10. %s' % safe_unicode(obj.detail_10) if obj.detail_10 else ''
        parameter['year'] = safe_unicode(obj.docDate.year - 1911)
        parameter['month'] = safe_unicode(obj.docDate.month)
        parameter['day'] = safe_unicode(obj.docDate.day)
        filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/official_document.docx'

        doc = DocxTemplate(filePath)
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile(title)


class GernalReport(Basic):
    def __call__(self):
        result = self.getAllStudent()
        context = self.context
        request = self.request
        mode = request.get('mode', 'horizon')
        course = context.getParentNode().Title()
        period = context.id

        if mode == 'horizon':
            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/horizon.docx'
            title = '%s_%s_橫式' %(course, period)
        elif mode == 'straight':
            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/straight.docx'
            title = '%s_%s_直式' %(course, period)
        elif mode == 'fire':
            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/fire.docx'
            title = '%s_%s_防火' %(course, period)
        elif mode == 'issue':
            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/issue.docx'
            title = '%s_%s_核發清冊' %(course, period)
        elif mode == 'emergency_grade':
            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/emergency_grade.docx'
            title = '%s_%s_成績冊' %(course, period)

        doc = DocxTemplate(filePath)

        start = self.getChineseBirthday(context.courseStart)
        end = self.getChineseBirthday(context.courseEnd)
        time = '訓練期間：%s~%s' %(start, end)
        trainingCenter = context.trainingCenter.to_object.title

        parameter = {}
        parameter['course'] = safe_unicode(course)
        parameter['period'] = safe_unicode(period)
        parameter['time'] = safe_unicode(time)
        parameter['trainingCenter'] = safe_unicode(trainingCenter)
        temp = []
        for item in result:
            obj = dict(item)
            obj['note'] = ''
            obj['birthday'] = self.getChineseBirthday(obj['birthday'])

            if course == '職業安全管理師' or course == '職業衛生管理師':
                license_date = self.getChineseBirthday(obj['license_date'])
                if obj['license_unit']:
                    obj['note'] += obj['license_unit'] + '\n'
                if license_date:
                    obj['note'] += license_date + '\n'
                if obj['license_code']:
                    obj['note'] += obj['license_code'] + '\n'
            elif re.match('防火', course):
                pass
            else:
                obj['note'] += obj['status'] + '\n'

            temp.append(obj)

        data = []
        count = 0
        length = len(result)
        while length > 0:
            data.append(temp[count * 8: (count + 1) * 8])
            length -= 8
            count += 1
        parameter['data'] = data
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile(title)


class FireSingInReport(Basic):
    def __call__(self):
        result = self.getAllStudent()
        context = self.context
        document = Document()
        sections = document.sections
        for section in sections:
            # change orientation to landscape
            section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        course = context.getParentNode().Title()
        period = context.id

        title = '中國勞工安全衛生管理學會%s訓練班第%s期學員簽到' %(course, period)
        style1 = document.styles.add_style('style1', 1)
        style1.font.size = Pt(24)
        document.add_paragraph(title.decode(), style = style1)

        count = 1
        table = document.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '編號'.decode()
        hdr_cells[1].text = '姓名'.decode()
        hdr_cells[2].text = '上午簽到'.decode()
        hdr_cells[3].text = '下午簽到'.decode()
        hdr_cells[4].text = '編號'.decode()
        hdr_cells[5].text = '姓名'.decode()
        hdr_cells[6].text = '上午簽到'.decode()
        hdr_cells[7].text = '下午簽到'.decode()

        for item in result:
            obj = dict(item)
            if count in [17, 33, 49, 65, 81, 97]:
                table = document.add_table(rows=1, cols=8)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '編號'.decode()
                hdr_cells[1].text = '姓名'.decode()
                hdr_cells[2].text = '上午簽到'.decode()
                hdr_cells[3].text = '下午簽到'.decode()
                hdr_cells[4].text = '編號'.decode()
                hdr_cells[5].text = '姓名'.decode()
                hdr_cells[6].text = '上午簽到'.decode()
                hdr_cells[7].text = '下午簽到'.decode()

            row_cells = table.add_row().cells
            row_cells[0].text = str(count)
            row_cells[1].text = obj['name'].decode()
            row_cells[4].text = str(count + 1)
            row_cells[5].text = obj['name'].decode()
            count += 2
            if count in [17, 33, 49, 65, 81, 97]:
                document.add_page_break()

        document.save('/tmp/temp.docx')

        return self.downloadFile(title)


class NormalSingInReport(Basic):
    def __call__(self):
        request = self.request
        result = self.getAllStudent()
        context = self.context
        document = Document()
        subject = request.get('subject')
        parameter = {}
        course = context.getParentNode().Title()
        period = context.id
        title =  '第%s期%s受訓學員簽到紀錄' %(period, course)
        trainingCenter = context.trainingCenter.to_object.title
        start = self.getChineseBirthday(context.courseStart)
        end = self.getChineseBirthday(context.courseEnd)
        time = '%s~%s' %(start, end)

        parameter['title'] = title
        parameter['trainingCenter'] = trainingCenter
        parameter['time'] = time
        parameter['subject'] = subject

        count = 1
        for item in result:
            parameter['seatNo%s' %count] = item[19]
            parameter['name%s' %count] = item[4]
            count += 1

        filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/normal_sing_in.docx'
        doc = DocxTemplate(filePath)
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile(title)



class DownloadGrade(Basic):
    def __call__(self):
        request = self.request
        result = self.getAllStudent()
        context = self.context
        document = Document()
        subject = request.get('subject')
        parameter = {}
        course = context.getParentNode().Title()
        period = context.id
        title =  '第%s期%s受訓學員成績冊' %(period, course)
        trainingCenter = context.trainingCenter.to_object.title
        start = self.getChineseBirthday(context.courseStart)
        end = self.getChineseBirthday(context.courseEnd)
        time = '%s~%s' %(start, end)

        parameter['title'] = title
        parameter['trainingCenter'] = trainingCenter
        parameter['time'] = time

        temp = []
        for item in result:
            obj = dict(item)
            obj['birthday'] = self.getChineseBirthday(obj['birthday'])
            temp.append(obj)

        data = []
        count = 0
        length = len(result)
        while length > 0:
            data.append(temp[count * 5: (count + 1) * 5])
            length -= 5
            count += 1
        parameter['data'] = data

        filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/grade.docx'
        doc = DocxTemplate(filePath)
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile(title)


class DownloadPicture(Basic):
    def __call__(self):
        request = self.request
        result = self.getAllStudent()
        context = self.context
        document = Document()
        subject = request.get('subject')
        parameter = {}
        course = context.getParentNode().Title()
        period = context.id
        title =  '第%s期%s' %(period, course)

        parameter['title'] = title

        data = []
        for item in result:
            seatNo = item[19]
            name = item[4]
            data.append('%s號 %s' %(seatNo, name))

        parameter['data'] = data
        filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/picture.docx'
        doc = DocxTemplate(filePath)
        doc.render(parameter)
        doc.save("/tmp/temp.docx")

        return self.downloadFile(title)


class CourseReport(Basic):
    template = ViewPageTemplateFile("template/course_report.pt")

    def __call__(self):
        request = self.request
        portal = api.portal.get()

        center = request.get('center')

        if center:
            periodContent = api.content.find(context=portal['mana_course'], trainingCenterId=center)
            trainingCenter = api.content.find(id=center)[0].getObject()
            parameter = {}
            data = []

            for item in periodContent:
                obj = item.getObject()

                regDeadline = obj.regDeadline.strftime('%Y-%m-%d') if obj.regDeadline else ''
                courseStart = obj.courseStart.strftime('%Y-%m-%d') if obj.courseStart else ''

                data.append({
                    'courseFee': obj.courseFee,
                    'duringTime': self.checkDuringTime(obj.duringTime),
                    'title': obj.getParentNode().Title(),
                    'regDeadline': regDeadline,
                    'courseStart': courseStart,
                    'classStatus': self.checkClassStatus(obj.classStatus)
                })

            document = Document()
            parameter['data'] = data
            parameter['trainingCenter'] = trainingCenter.title
            parameter['phone'] = trainingCenter.phone
            parameter['fax'] = trainingCenter.fax
            parameter['address'] = trainingCenter.address
            parameter['now'] = datetime.datetime.now().strftime('%Y-%m-%d')

            filePath = '/home/andy/cshm/zeocluster/src/cshm.content/src/cshm/content/browser/static/course_report.docx'
            doc = DocxTemplate(filePath)
            doc.render(parameter)
            doc.save("/tmp/temp.docx")

            return self.downloadFile('中國勞工安全衛生管理學會%s近期開課一覽表' %trainingCenter.title)

        else:
            centerList = []

            content = api.content.find(context=portal['resource']['training_center'], depth=1)
            for item in content:
                obj = item.getObject()
                centerList.append({
                    'title': obj.title,
                    'code':obj.id,
                })
            self.centerList = centerList
            return self.template()

    def checkDuringTime(self, duringTime):
        if duringTime == 'notYat':
            return '時段未定'
        elif duringTime == 'inDay':
            return '日間'
        elif duringTime == 'inEvening':
            return '夜間'
        elif duringTime == 'inWeekend':
            return '假日'
        elif duringTime == 'inWeekendEvening':
            return '周末夜間'
        elif duringTime == 'complex':
            return '混合時段'
        elif duringTime == 'phone':
            return '電話'

    def checkClassStatus(self, status):
        if status == 'willStart':
            return '確定開課'
        elif status == 'fullCanAlt':
            return '額滿候補'
        elif status == 'planed':
            return '預定開課'
        elif status == 'registerFirst':
            return '請先報名'
        elif status == 'altFull':
            return '報名額滿'
